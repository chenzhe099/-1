"""
生成嵌入流程图的商业计划书 Word 文档 v2
在对应章节自动插入 PNG 流程图
"""
import re, os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE = r'd:\zhuomian\gjcx'
IMG_DIR = os.path.join(BASE, 'assets', 'flowcharts')

def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_styled_table(doc, headers, rows, header_color='1A5276'):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h); run.bold = True; run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255,255,255); run.font.name = 'Microsoft YaHei'
        set_cell_shading(cell, header_color)
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r+1].cells[c]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c>0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val)); run.font.size = Pt(9); run.font.name = 'Microsoft YaHei'
            if r%2==0: set_cell_shading(cell, 'F8F9FA')
    doc.add_paragraph()
    return table

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs: run.font.name = 'Microsoft YaHei'
    return h

def add_para(doc, text, bold=False, size=10.5, align=None, spacing_after=6):
    p = doc.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text); run.font.size = Pt(size); run.font.name = 'Microsoft YaHei'
    if bold: run.bold = True
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet'); p.clear()
    run = p.add_run(text); run.font.size = Pt(10); run.font.name = 'Microsoft YaHei'

def insert_image(doc, filename, width_cm=15):
    """居中插入流程图"""
    path = os.path.join(IMG_DIR, filename)
    if not os.path.exists(path):
        add_para(doc, f'[图片待插入: {filename}]', size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    # 图片编号
    add_para(doc, filename.replace('.png','').replace('_',' '), size=7.5,
             align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=12)

def clean_text(text):
    """清理 markdown 格式残留"""
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'✅', '[完成]', text)
    text = re.sub(r'🔲', '[待开发]', text)
    text = re.sub(r'🔴', '[高]', text)
    text = re.sub(r'🟡', '[中]', text)
    text = re.sub(r'🟢', '[低]', text)
    return text

def build_doc(md_path, docx_path):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21); section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54); section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.6); section.right_margin = Cm(2.6)
    style = doc.styles['Normal']; style.font.name = 'Microsoft YaHei'; style.font.size = Pt(10.5)

    # ========== 封面 ==========
    for _ in range(5): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('智 农 云'); run.bold = True; run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(26,82,118); run.font.name = 'Microsoft YaHei'

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('SmartAgri Cloud'); run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100,100,100); run.font.name = 'Arial'

    doc.add_paragraph()
    for title in ['智慧农业管理系统', '商 业 计 划 书']:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title); run.bold = True; run.font.size = Pt(26)
        run.font.color.rgb = RGBColor(26,82,118); run.font.name = 'Microsoft YaHei'

    doc.add_paragraph()
    for label, value in [('版本：','V1.0'),('日期：','2024年1月'),('密级：','商业机密'),('状态：','天使轮融资')]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label); run.font.size = Pt(11); run.font.name = 'Microsoft YaHei'
        run = p.add_run(value); run.bold = True; run.font.size = Pt(11); run.font.name = 'Microsoft YaHei'

    for _ in range(3): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('智农云科技（筹）'); run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128,128,128); run.font.name = 'Microsoft YaHei'

    # ========== 目录 ==========
    doc.add_page_break()
    add_heading(doc, '目  录', level=1)
    doc.add_paragraph()
    toc = [('一、','执行摘要'),('二、','公司概述'),('三、','产品与技术'),('四、','市场分析'),
           ('五、','竞争分析'),('六、','商业模式'),('七、','市场推广策略'),('八、','财务预测'),
           ('九、','团队与运营'),('十、','风险评估与应对'),('十一、','发展路线图'),('十二、','融资需求'),
           ('附录A：','产品功能清单'),('附录B：','技术架构详情')]
    for num,title in toc:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8); p.paragraph_format.line_spacing = 1.8
        run = p.add_run(f'{num}{title}'); run.font.size = Pt(11); run.font.name = 'Microsoft YaHei'

    # ========== 正文解析 ==========
    doc.add_page_break()

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    start = content.find('## 一、执行摘要')
    if start < 0: start = 0
    body = content[start:]

    # 流程图插入映射: 在遇到特定章节标题后插入对应图片
    img_triggers = {
        '### 3.1 产品架构': '01_product_architecture.png',
        '### 3.3 技术栈': '07_tech_stack.png',
        '### 5.1 竞争格局': '06_competitive_matrix.png',
        '### 6.1 收入来源': '05_business_model.png',
        '### 7.1 获客渠道': '02_acquisition_channels.png',
        '### 8.1 3 年损益预测': '08_financial_projection.png',
        '### 9.1 组织架构': '03_org_chart.png',
        '## 十、风险评估与应对': '09_risk_matrix.png',
        '## 十一、发展路线图': '04_roadmap.png',
    }

    lines = body.split('\n')
    i = 0
    in_table = False; table_headers = []; table_rows = []
    in_code = False

    while i < len(lines):
        line = lines[i]

        # 代码块：跳过全部
        if line.strip().startswith('```'):
            in_code = not in_code
            i += 1; continue
        if in_code:
            i += 1; continue

        # 图片插入触发
        if i < len(lines) - 1:
            combined = line.strip()
            h2_or_h3 = (combined.startswith('## ') or combined.startswith('### ')) and combined[3:].strip()
            for trigger, img_file in img_triggers.items():
                if h2_or_h3 and h2_or_h3 in trigger:
                    # 先不插入，等标题渲染后再插入
                    pass

        if line.strip() == '---':
            doc.add_paragraph(); i += 1; continue

        # H2
        if line.startswith('## ') and not line.startswith('### '):
            text = clean_text(line[3:].strip())
            add_heading(doc, text, level=1)
            # 检查是否需要插入图片
            for trigger, img_file in img_triggers.items():
                if trigger.startswith('## ') and trigger[3:] in line:
                    add_para(doc, '', spacing_after=2)
                    insert_image(doc, img_file)
            i += 1; continue

        # H3
        if line.startswith('### '):
            text = clean_text(line[4:].strip())
            add_heading(doc, text, level=2)
            for trigger, img_file in img_triggers.items():
                if trigger.startswith('### ') and trigger[4:] in line:
                    add_para(doc, '', spacing_after=2)
                    insert_image(doc, img_file)
            i += 1; continue

        # H4
        if line.startswith('#### '):
            add_heading(doc, clean_text(line[5:].strip()), level=3)
            i += 1; continue

        # 表格处理
        if line.strip().startswith('|') and line.strip().endswith('|'):
            if not in_table:
                in_table = True
                table_headers = [clean_text(c.strip()) for c in line.split('|')[1:-1]]
            elif re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                pass  # separator
            else:
                table_rows.append([clean_text(c.strip()) for c in line.split('|')[1:-1]])
            if i+1 >= len(lines) or not lines[i+1].strip().startswith('|'):
                if table_headers and table_rows:
                    add_styled_table(doc, table_headers, table_rows)
                table_headers = []; table_rows = []; in_table = False
            i += 1; continue

        # 空行
        if not line.strip():
            i += 1; continue

        # 引用
        if line.strip().startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1); p.paragraph_format.space_after = Pt(4)
            run = p.add_run(clean_text(line.strip()[2:])); run.font.size = Pt(9)
            run.font.italic = True; run.font.color.rgb = RGBColor(100,100,100)
            run.font.name = 'Microsoft YaHei'
            i += 1; continue

        # 列表
        if re.match(r'^[\-\*\d]+\.?\s', line.strip()):
            text = re.sub(r'^[\-\*\d]+\.?\s', '', line.strip())
            text = clean_text(text)
            if text: add_bullet(doc, text)
            i += 1; continue

        # 普通段落
        text = clean_text(line.strip())
        if text:
            # 跳过纯加粗标记残留
            if text in ['**', '---', '...'] or text.startswith('┌') or text.startswith('│') or text.startswith('└'):
                i += 1; continue
            add_para(doc, text)
        i += 1

    # ========== 页眉页脚 ==========
    footer = section.footer; footer.is_linked_to_previous = False
    p = footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('— '); run.font.size = Pt(8); run.font.color.rgb = RGBColor(150,150,150)
    fld1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    fld2 = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    fld3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run._r.append(fld1); run._r.append(fld2); run._r.append(fld3)
    run = p.add_run(' —'); run.font.size = Pt(8); run.font.color.rgb = RGBColor(150,150,150)

    header = section.header; header.is_linked_to_previous = False
    p = header.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('智农云 — 商业计划书  |  商业机密'); run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150,150,150); run.font.name = 'Microsoft YaHei'

    doc.save(docx_path)
    print(f'OK Document saved: {docx_path}')
    print(f'File size: {os.path.getsize(docx_path)/1024:.1f} KB')


if __name__ == '__main__':
    md_path = os.path.join(BASE, '商业计划书.md')
    docx_path = os.path.join(BASE, '智农云_商业计划书_V1.0.docx')
    build_doc(md_path, docx_path)
