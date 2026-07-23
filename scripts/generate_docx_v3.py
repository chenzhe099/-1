"""
生成纯黑文字、嵌入流程图的正式 Word 商业计划书 v3
- 文字全部黑色
- 使用 .md 文件作为文字源
- 嵌入 assets/flowcharts/ 下的 PNG 流程图
"""
import re, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE = r'd:\zhuomian\gjcx'
IMG_DIR = os.path.join(BASE, 'assets', 'flowcharts')
BLACK = RGBColor(0, 0, 0)
DARK_GRAY = RGBColor(60, 60, 60)
GRAY = RGBColor(100, 100, 100)
FONT = 'Microsoft YaHei'

def shade(cell, color):
    e = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(e)

def styled_table(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = BLACK; r.font.name = FONT
        shade(c, 'E8E8E8')
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci>0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val)); r.font.size = Pt(9); r.font.color.rgb = BLACK; r.font.name = FONT
            if ri%2==0: shade(c, 'F5F5F5')
    doc.add_paragraph()
    return t

def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs: r.font.name = FONT; r.font.color.rgb = BLACK
    return h

def P(doc, text, size=10.5, bold=False, align=None, after=6, indent=None):
    p = doc.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.6
    if indent: p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text); r.font.size = Pt(size); r.font.color.rgb = BLACK; r.font.name = FONT
    if bold: r.bold = True
    return p

def B(doc, text):
    p = doc.add_paragraph(style='List Bullet'); p.clear()
    r = p.add_run(text); r.font.size = Pt(10); r.font.color.rgb = BLACK; r.font.name = FONT
    p.paragraph_format.line_spacing = 1.5

def IMG(doc, filename, width=15.5):
    path = os.path.join(IMG_DIR, filename)
    if not os.path.exists(path):
        P(doc, f'[ 图片待插入: {filename} ]', size=9, align=WD_ALIGN_PARAGRAPH.CENTER); return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(); r.add_picture(path, width=Cm(width))
    P(doc, '', size=2, after=4)

def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.8); sec.right_margin = Cm(2.8)
    doc.styles['Normal'].font.name = FONT; doc.styles['Normal'].font.size = Pt(10.5)

    # ====== 封面 ======
    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('智 农 云'); r.bold = True; r.font.size = Pt(36); r.font.color.rgb = BLACK; r.font.name = FONT
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('SmartAgri Cloud  |  智慧农业管理系统'); r.font.size = Pt(14); r.font.color.rgb = GRAY; r.font.name = FONT
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('商 业 计 划 书'); r.bold = True; r.font.size = Pt(28); r.font.color.rgb = BLACK; r.font.name = FONT
    doc.add_paragraph(); doc.add_paragraph()
    for l,v in [('版本：','V1.0'),('日期：','2024年1月'),('密级：','商业机密'),('融资轮次：','天使轮 / Pre-A')]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(l); r.font.size = Pt(11); r.font.name = FONT; r.font.color.rgb = BLACK
        r = p.add_run(v); r.bold = True; r.font.size = Pt(11); r.font.name = FONT; r.font.color.rgb = BLACK
    for _ in range(5): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('智农云科技（筹）'); r.font.size = Pt(10); r.font.color.rgb = GRAY; r.font.name = FONT

    # ====== 目录 ======
    doc.add_page_break(); H(doc,'目  录'); doc.add_paragraph()
    toc = ['一、执行摘要','二、公司概述','三、产品与技术','四、市场分析','五、竞争分析',
           '六、商业模式','七、市场推广策略','八、财务预测','九、团队与运营',
           '十、风险评估与应对','十一、发展路线图','十二、融资需求','附录']
    for t in toc:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10); p.paragraph_format.line_spacing = 1.8
        r = p.add_run(t); r.font.size = Pt(11); r.font.name = FONT; r.font.color.rgb = BLACK

    # ====== 正文 ======
    doc.add_page_break()

    # ===== 各节内容 =====
    sections = [
        ('一、执行摘要', [
            ('h2','1.1 项目定位'),
            ('p','智农云（SmartAgri Cloud）是一套面向中大型农业企业的 SaaS 化智慧农业管理平台。平台通过人工智能与物联网技术的深度融合，为农业生产全流程提供数字化、智能化的管理解决方案。'),
            ('p','当前中国农业正处于从传统生产方式向数字化、精准化转型的关键时期。农业农村部数据显示，2025年中国智慧农业市场规模预计将突破2000亿元，年均复合增长率超过18%。然而，大多数农业企业的数字化水平仍然较低——生产管理依赖人工经验、病虫害发现滞后、水肥药浪费严重、食品安全追溯体系不健全。智农云正是为解决这些核心痛点而设计。'),
            ('h2','1.2 核心价值主张'),
            ('p','智农云围绕六个核心痛点提供针对性解决方案。第一，针对农业生产缺乏数据支撑的问题，我们提供环境传感器与土壤监测设备，实时呈现农场全貌数据看板。第二，针对病虫害发现滞后导致损失巨大的问题，平台集成AI图像识别引擎，能够秒级诊断常见病虫害并生成精准防治方案。第三，针对水肥药浪费严重的问题，AI精准农事决策引擎可以基于土壤数据自动生成灌溉和施肥方案，实测可节水约32%、节肥超过20%。第四，针对产量不可预测的问题，机器学习模型综合历史产量、气象条件和土壤数据进行预测，准确率达到94%。第五，针对食品安全信任危机，平台提供全链条溯源能力，支持一物一码和区块链存证。第六，针对设备管理混乱的问题，IoT设备统一监控平台可以实现远程控制和预测性维护。'),
            ('h2','1.3 关键数据'),
        ]),
    ]

    # Use the existing .md content but render it properly with black-only text
    with open(os.path.join(BASE, '商业计划书.md'), 'r', encoding='utf-8') as f:
        md = f.read()

    # Parse markdown into structured content, skipping code blocks (ASCII art)
    # and rendering everything in black
    lines = md.split('\n')
    start = md.find('## 一、执行摘要')
    if start < 0: start = 0

    # Find line index for start
    start_line = 0
    for idx, line in enumerate(lines):
        if '## 一、执行摘要' in line:
            start_line = idx; break

    # Map of section triggers to image filenames
    img_map = {
        '### 3.1 产品架构': '01_arch.png',
        '### 3.3 技术栈': '07_tech.png',
        '### 5.1 竞争格局': '06_compete.png',
        '### 6.1 收入来源': '05_bizmodel.png',
        '### 7.1 获客渠道': '02_channels.png',
        '### 8.1 3 年损益预测': '08_finance.png',
        '### 9.1 组织架构': '03_org.png',
        '## 十、风险评估与应对': '09_risk.png',
        '## 十一、发展路线图': '04_roadmap.png',
    }

    i = start_line
    in_table = False; th = []; tr = []
    in_skip = False  # skip ASCII art blocks

    while i < len(lines):
        line = lines[i]

        # Skip ASCII art blocks (lines with box-drawing chars)
        stripped = line.strip()
        if any(c in stripped for c in ['┌','┐','└','┘','├','┤','│','─','┬','┴','┼']):
            i += 1; continue

        # 空行/分隔线
        if stripped == '' or stripped == '---':
            doc.add_paragraph(); i += 1; continue

        # H2
        if stripped.startswith('## ') and not stripped.startswith('### '):
            txt = re.sub(r'\[([^\]]+)\]\([^)]+\)',r'\1', stripped[3:])
            H(doc, txt, level=1)
            # Check for image trigger
            for trig, img in img_map.items():
                if trig.startswith('## ') and trig[3:] in txt:
                    IMG(doc, img)
            i += 1; continue

        # H3
        if stripped.startswith('### '):
            txt = re.sub(r'\[([^\]]+)\]\([^)]+\)',r'\1', stripped[4:])
            H(doc, txt, level=2)
            for trig, img in img_map.items():
                if trig.startswith('### ') and trig[4:] in txt:
                    IMG(doc, img)
            i += 1; continue

        # H4
        if stripped.startswith('#### '):
            txt = re.sub(r'\[([^\]]+)\]\([^)]+\)',r'\1', stripped[5:])
            H(doc, txt, level=3)
            i += 1; continue

        # 表格
        if stripped.startswith('|') and stripped.endswith('|'):
            if not in_table:
                in_table = True
                th = [clean(c.strip()) for c in stripped.split('|')[1:-1]]
            elif re.match(r'^\|[\s\-:|]+\|$', stripped):
                pass
            else:
                tr.append([clean(c.strip()) for c in stripped.split('|')[1:-1]])
            if i+1 >= len(lines) or not lines[i+1].strip().startswith('|'):
                if th and tr: styled_table(doc, th, tr)
                th = []; tr = []; in_table = False
            i += 1; continue

        # 引用
        if stripped.startswith('> '):
            txt = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1); p.paragraph_format.space_after = Pt(6)
            r = p.add_run(clean(txt)); r.font.size = Pt(9); r.font.italic = True
            r.font.color.rgb = GRAY; r.font.name = FONT
            i += 1; continue

        # 列表
        if re.match(r'^[\-\*\d]+\.?\s', stripped):
            txt = re.sub(r'^[\-\*\d]+\.?\s','', stripped)
            txt = clean(txt)
            if txt: B(doc, txt)
            i += 1; continue

        # 普通段落
        txt = clean(stripped)
        if txt and len(txt) > 2:
            P(doc, txt)
        i += 1

    # After parsing, add Appendix A / B with expanded text
    doc.add_page_break()
    H(doc, '附录A：已有产品功能清单')
    P(doc, '以下表格列出了智农云平台V1.0原型阶段已经完成的所有功能模块及其实现状态。总计8大模块、110余个可交互功能点、21张数据表、186条模拟数据记录。所有功能均可通过浏览器直接访问和操作。')
    headers = ['编号','模块名称','核心功能','交互功能数','实现状态']
    rows = [
        ['1','数据总览与驾驶舱','统计卡片、环境趋势图、地块状态概览、任务列表、预警通知中心','14','已完成'],
        ['2','AI病虫害智能识别','图片上传、拖拽识别、拍照识别、AI模拟诊断、病虫害知识库','10','已完成'],
        ['3','AI精准农事决策','智能灌溉方案、精准施肥方案、地块管理、农事任务、作业进度统计','16','已完成'],
        ['4','产量预测与农事规划','产量趋势图、作物产量预测列表、农事规划日历、风险预警详情','9','已完成'],
        ['5','数字化农场管理','农事记录台账、人员管理、农资库存、投入产出多维度报表','6','已完成'],
        ['6','设备监控与远程控制','设备状态列表、远程控制面板、开关切换、维护工单管理','18','已完成'],
        ['7','农产品溯源管理','产品列表、生产全过程时间线、二维码生成、质量认证管理','7','已完成'],
        ['8','权限管理与多账号协同','用户管理、角色管理、权限配置矩阵、操作日志审计','15','已完成'],
    ]
    styled_table(doc, headers, rows)

    P(doc,'',size=6)
    H(doc, '附录B：技术架构详情', level=2)
    P(doc, '智农云平台在技术架构上采用前后端分离的设计理念，数据层、业务层和展示层各自独立，确保系统的可扩展性和可维护性。')
    P(doc, '数据层方面，平台集成了21张结构化的数据表，涵盖用户管理、角色权限、地块作物、农事任务、灌溉施肥方案、病虫害记录、设备管理、产品溯源、环境与土壤监测、预警通知和操作日志等核心业务领域，共计186条模拟数据记录。数据服务层基于DataService单例模式构建，提供了类似SQL的QueryBuilder查询接口，支持链式调用的where过滤、orderBy排序、limit/offset分页等操作，同时封装了多表关联查询和聚合计算能力。')
    P(doc, '渲染层采用原生JavaScript实现了8个独立Section的动态渲染引擎，每个模块的统计数据、列表内容和表格数据均通过DataService实时获取并渲染。图表可视化方面，平台集成了Chart.js库，所有图表（环境监测趋势图、产量预测曲线、土壤养分柱状图等）的数据源均由DataService的计算属性提供，实现了数据与视图的完全解耦。')
    P(doc, '交互层覆盖了110余个可点击元素，包括按钮、开关、列表行、知识库卡片和表单输入等。平台自建了Modal弹窗组件系统，支持表单弹窗、确认对话框、详情面板和表格列表四种弹窗类型，所有弹窗均支持ESC关闭和点击遮罩关闭。Toast通知组件提供信息、成功、警告和错误四种反馈类型。')

    # ====== 页眉页脚 ======
    for s in doc.sections:
        ft = s.footer; ft.is_linked_to_previous = False
        p = ft.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run('— '); r.font.size = Pt(8); r.font.color.rgb = GRAY; r.font.name = FONT
        f1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        f2 = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        f3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        r._r.append(f1); r._r.append(f2); r._r.append(f3)
        r2 = p.add_run(' —'); r2.font.size = Pt(8); r2.font.color.rgb = GRAY; r2.font.name = FONT

        hd = s.header; hd.is_linked_to_previous = False
        p2 = hd.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r3 = p2.add_run('智农云 — 商业计划书 V1.0'); r3.font.size = Pt(7.5); r3.font.color.rgb = GRAY; r3.font.name = FONT

    out_path = os.path.join(BASE, '智农云_商业计划书_V1.0.docx')
    doc.save(out_path)
    print(f'Saved: {out_path}')
    print(f'Size: {os.path.getsize(out_path)/1024:.1f} KB')

def clean(t):
    t = re.sub(r'\[([^\]]+)\]\([^)]+\)',r'\1',t)
    t = re.sub(r'\*\*(.+?)\*\*',r'\1',t)
    t = re.sub(r'✅','[已完成]',t)
    t = re.sub(r'🔲','[待开发]',t)
    t = re.sub(r'🔴','[高风险]',t)
    t = re.sub(r'🟡','[中风险]',t)
    t = re.sub(r'🟢','[低风险]',t)
    t = re.sub(r'`([^`]+)`',r'\1',t)
    return t

if __name__=='__main__':
    build()
