"""
生成正式格式的商业计划书 Word 文档
使用 python-docx 实现专业排版
"""
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

def set_cell_shading(cell, color):
    """设置单元格底色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val","single")}" '
            f'w:sz="{val.get("sz","4")}" w:color="{val.get("color","auto")}" '
            f'w:space="0"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_styled_table(doc, headers, rows, col_widths=None, header_color='1A5276'):
    """添加格式化的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        set_cell_shading(cell, header_color)

    # 数据行
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            # 交替行背景色
            if r % 2 == 0:
                set_cell_shading(cell, 'F8F9FA')

    # 列宽
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing
    return table

def add_heading(doc, text, level=1):
    """添加标题，使用微软雅黑"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(doc, text, bold=False, size=10.5, align=None, color=None, spacing_after=6):
    """添加段落"""
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(spacing_after)
    pf.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(doc, text, level=0):
    """添加项目符号"""
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    return p

def add_code_block(doc, text):
    """添加代码块/ASCII图"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.0
    run = p.add_run(text)
    run.font.size = Pt(7.5)
    run.font.name = 'Consolas'
    run.font.color.rgb = RGBColor(80, 80, 80)
    return p

def parse_markdown_to_docx(md_path, docx_path):
    doc = Document()

    # ====== 页面设置 ======
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # ====== 设置默认字体 ======
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ====== 封面 ======
    # 空行压缩封面内容到上半部分
    for _ in range(6):
        doc.add_paragraph()

    # 公司名
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('智 农 云')
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(26, 82, 118)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 副标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('SmartAgri Cloud')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.font.name = 'Arial'

    doc.add_paragraph()

    # 主标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('智慧农业管理系统')
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(26, 82, 118)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('商 业 计 划 书')
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(26, 82, 118)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 分隔线
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 30)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(26, 82, 118)

    doc.add_paragraph()

    # 版本信息
    info_lines = [
        ('版本：', 'V1.0'),
        ('日期：', '2024年1月'),
        ('密级：', '商业机密'),
        ('状态：', '天使轮融资')
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run = p.add_run(value)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 封面页脚
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('智农云科技（筹）')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ====== 分页：目录页 ======
    doc.add_page_break()

    add_heading(doc, '目  录', level=1)
    doc.add_paragraph()

    toc_items = [
        ('一、', '执行摘要', '3'),
        ('二、', '公司概述', '5'),
        ('三、', '产品与技术', '6'),
        ('四、', '市场分析', '10'),
        ('五、', '竞争分析', '12'),
        ('六、', '商业模式', '13'),
        ('七、', '市场推广策略', '15'),
        ('八、', '财务预测', '16'),
        ('九、', '团队与运营', '18'),
        ('十、', '风险评估与应对', '19'),
        ('十一、', '发展路线图', '20'),
        ('十二、', '融资需求', '21'),
        ('附录A：', '已有产品功能清单', '23'),
        ('附录B：', '技术架构详情', '23'),
    ]
    for num, title, page in toc_items:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(8)
        pf.line_spacing = 1.8
        run = p.add_run(f'{num}{title}')
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        # 页码用制表符右对齐
        run = p.add_run(f'  ····················  {page}')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(150, 150, 150)

    # ====== 正文部分 ======
    doc.add_page_break()

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # 跳过封面和目录相关内容，从"## 一、执行摘要"开始
    start_marker = '## 一、执行摘要'
    start_idx = content.find(start_marker)
    if start_idx < 0:
        start_idx = 0
    body = content[start_idx:]

    lines = body.split('\n')
    i = 0
    in_table = False
    table_rows = []
    table_headers = []
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # 代码块
        if line.strip().startswith('```'):
            if in_code_block:
                # 结束代码块
                for cl in code_lines:
                    add_code_block(doc, cl)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 分隔线
        if line.strip() == '---':
            doc.add_paragraph()
            i += 1
            continue

        # H2: 主章节
        if line.startswith('## ') and not line.startswith('### '):
            text = line[3:].strip()
            # 清理锚点链接残留
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            add_heading(doc, text, level=1)
            i += 1
            continue

        # H3: 小节
        if line.startswith('### '):
            text = line[4:].strip()
            add_heading(doc, text, level=2)
            i += 1
            continue

        # H4: 子小节
        if line.startswith('#### '):
            text = line[5:].strip()
            add_heading(doc, text, level=3)
            i += 1
            continue

        # 表格
        if line.strip().startswith('|') and line.strip().endswith('|'):
            if not in_table:
                in_table = True
                table_headers = [c.strip() for c in line.split('|')[1:-1]]
            elif re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                # 分隔线，跳过
                pass
            else:
                row = [c.strip() for c in line.split('|')[1:-1]]
                table_rows.append(row)

            # 检查下一行是否还是表格
            if i + 1 < len(lines) and lines[i + 1].strip().startswith('|'):
                i += 1
                continue
            else:
                # 表格结束，渲染
                if table_headers and table_rows:
                    add_styled_table(doc, table_headers, table_rows)
                table_headers = []
                table_rows = []
                in_table = False
            i += 1
            continue

        # 空行
        if not line.strip():
            i += 1
            continue

        # 加粗文本行（如 **智农云**...）
        if line.strip().startswith('**') and '**' in line[2:]:
            text = line.strip()
            # 提取加粗内容
            bold_parts = re.findall(r'\*\*(.+?)\*\*', text)
            remaining = re.sub(r'\*\*(.+?)\*\*', '{}', text)
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_after = Pt(6)
            pf.line_spacing = 1.5

            parts = remaining.split('{}')
            for j, part in enumerate(parts):
                if part:
                    run = p.add_run(part)
                    run.font.size = Pt(10.5)
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                if j < len(bold_parts):
                    run = p.add_run(bold_parts[j])
                    run.bold = True
                    run.font.size = Pt(10.5)
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            i += 1
            continue

        # > 引用
        if line.strip().startswith('> '):
            text = line.strip()[2:]
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Cm(1)
            pf.space_after = Pt(4)
            run = p.add_run(text)
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            i += 1
            continue

        # 列表项 (- )
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            # 清理 markdown 加粗标记
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            add_bullet(doc, text)
            i += 1
            continue

        # 有序列表 (1. )
        if re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            add_bullet(doc, text)
            i += 1
            continue

        # 普通段落
        text = line.strip()
        # 清理格式标记
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        if text:
            add_para(doc, text)
        i += 1

    # ====== 页眉页脚 ======
    # 添加页脚（页码）
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('— ')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)
    # 添加自动页码
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run = p.add_run(' —')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)

    # 添加页眉
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('智农云 — 商业计划书  |  商业机密')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ====== 保存 ======
    doc.save(docx_path)
    print(f'文档已生成: {docx_path}')
    print(f'文件大小: {os.path.getsize(docx_path) / 1024:.1f} KB')


if __name__ == '__main__':
    base = r'd:\zhuomian\gjcx'
    md_path = os.path.join(base, '商业计划书.md')
    docx_path = os.path.join(base, '智农云_商业计划书_V1.0.docx')
    parse_markdown_to_docx(md_path, docx_path)
